from docx import Document
from docx.shared import Inches
from math import *
from Data import *
import win32com.client as win32
import os, sys

# creating a new document

document = Document()

document.add_heading('Game Summary', 0)
document.add_paragraph("The words given to you were :")

row = str()

for i in range(5):
    for j in range(5):
        if((i * 5 + j + 1) // 10 == 0):
            row = row + "  "
        row = row + str(i * 5 + j + 1) + ". " + str(anagselec[(i * 5 + j)]) + "\t"
    paragraph = document.add_paragraph(row)
    paragraph_format = paragraph.paragraph_format
    tab_stops = paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(0.97))      # <- Dont loop this
    tab_stops.add_tab_stop(Inches(2*0.97))
    tab_stops.add_tab_stop(Inches(3*0.97))
    tab_stops.add_tab_stop(Inches(4*0.97))
    row = ""

document.add_paragraph("\nFollowing anaglinks were present :")

for lnk in ans_copy:
    for anag in lnk:
        row = row + " - " + anag
    document.add_paragraph(row)
    row = ""

document.add_paragraph("\nYour current status : " + stat + ".")

if( stat == 'Won' ):
    document.add_paragraph("\nYou scored : " + str(Score) + ",  Which includes a time bonus : " + str(floor((60 - timer) * 0.5)))

document.save('summary.docx')

filepath = os.path.join(os.path.dirname(sys.argv[0]), 'summary.docx')

wordApp = win32.Dispatch('Word.Application') #create a word application object
wordApp.Visible = False # hide the word application
doc = wordApp.Documents.Open(filepath)

#Adding hyperlink

for wrd in anagselec:
    if wordApp.Selection.Find.Execute(FindText = wrd):
        doc.Hyperlinks.Add(Anchor = wordApp.Selection.Range,  Address=("https://en.wiktionary.org/wiki/"+wrd))

doc.Save()
doc.Close()

wordApp.Visible = True # unhide the word application
doc = wordApp.Documents.Open(filepath, False, False, True)

foregroundWindow("summary")



